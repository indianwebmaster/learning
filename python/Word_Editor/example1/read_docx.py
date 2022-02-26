import docx

doc = docx.Document("NewITACommunityServiceLetterTEMPLATE.docx")
print(len(doc.paragraphs))

for paragraph in doc.paragraphs:
    if 'keyTotalHours' in paragraph.text:
        newParagraph = paragraph.text.replace("keyTotalHours","10:00")
    else:
        newParagraph = paragraph.text
    print(newParagraph)
